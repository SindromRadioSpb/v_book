"""DOCX extractor."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: Path) -> str:
    """
    Extract text from DOCX file.

    Extracts all paragraphs and table cells, preserving document structure.
    Returns empty string if python-docx is not available.
    """
    logger.info(f"Extracting text from DOCX: {file_path}")

    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise ImportError("python-docx library is required for DOCX extraction")

    try:
        doc = Document(file_path)
        paragraphs = []

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        result = "\n".join(paragraphs)
        logger.info(f"Extracted {len(result)} characters from DOCX")
        return result

    except Exception as e:
        logger.exception(f"Failed to extract text from DOCX: {file_path}")
        raise RuntimeError(f"DOCX extraction failed: {e}")
